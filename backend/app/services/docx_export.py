"""Generación de documentos DOCX para exportar borradores de propuesta técnica."""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from app.models.analisis_ia import BorradorPropuesta


_COLOR_PRIMARIO = RGBColor(0x1E, 0x40, 0xAF)  # blue-800


def generar_docx_borrador(borrador: "BorradorPropuesta") -> io.BytesIO:
    """Genera un documento DOCX a partir del borrador de propuesta técnica.

    Devuelve un buffer listo para ser enviado como descarga HTTP.
    """
    doc = Document()
    _configurar_estilos(doc)

    # Título principal
    titulo = borrador.titulo or "Propuesta Técnica"
    titulo_par = doc.add_heading(titulo, level=0)
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_par.runs:
        run.font.color.rgb = _COLOR_PRIMARIO

    doc.add_paragraph("")  # espacio

    # Secciones del cuerpo
    secciones: list[Any] = borrador.secciones or []
    for seccion in secciones:
        nombre = seccion.get("titulo", "") if isinstance(seccion, dict) else ""
        contenido = seccion.get("contenido", "") if isinstance(seccion, dict) else ""
        if nombre:
            doc.add_heading(nombre, level=1)
        if contenido:
            par = doc.add_paragraph(contenido)
            par.paragraph_format.space_after = Pt(8)

    # Documentos a preparar
    docs_pendientes: list[Any] = borrador.documentos_pendientes or []
    if docs_pendientes:
        doc.add_heading("Documentos a preparar", level=1)
        for item in docs_pendientes:
            doc.add_paragraph(str(item), style="List Bullet")

    # Notas de revisión
    notas: list[Any] = borrador.notas_revision or []
    if notas:
        doc.add_heading("Notas de revisión", level=1)
        for nota in notas:
            doc.add_paragraph(str(nota), style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _configurar_estilos(doc: Document) -> None:
    """Ajusta tamaño de fuente base del documento."""
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
